import os
from docx import Document
from win32com import client as wc
import xlwt


class Solution:
    def getdirfiles(self, dirname):
        self.cleartempfiles(dirname)
        dicts = {}
        for root, dirs, files in os.walk(dirname):
            for file in files:
                # print(os.path.join(root, file))
                filename = os.path.join(root, file)
                shortname, extension = os.path.splitext(file)
                excelfilename = shortname + '.xls'
                excelfilename = os.path.join(root, excelfilename)
                print('excelfilename:',excelfilename)
                if os.path.exists(excelfilename):
                    dicts[shortname] = 0
                if shortname not in dicts:
                    if not filename.startswith('~$'):
                        if filename.endswith('.doc') or filename.endswith('.docx'):
                            print('process:', filename)
                            # tablelength = []
                            # returnedvalue, tablelength = self.getdocxexcel(filename)
                            # dict[filename] = tablelength
                            dicts[shortname] = self.getdocxexcel(filename)
        print('处理文件如下：')
        filename1 = dirname + '\\已处理文件.txt'
        filename2 = dirname + '\\非标准、未处理文件.txt'
        f1 = open(filename1, 'w')
        for k, v in dicts.items():
            if v == 0:
                print(k, v)
                tempstr = k + '\n'
                f1.write(tempstr)
        f1.close()
        print('未处理，非标准格式文件如下：')
        f2 = open(filename2, 'w')
        for k, v in dicts.items():
            if v == 1:
                print(k, v)
                tempstr = k + '\n'
                f2.write(tempstr)
        f2.close()
        self.cleartempfiles(dirname)
        return

    def cleartempfiles(self, dirname):
        #清理临时文件
        count = 0
        for root, dirs, files in os.walk(dirname):
            for file in files:
                filename = os.path.join(root, file)
                if file.startswith('~$'):
                    count += 1
                    os.remove(filename)
        print('共清理文件数量：', count)

    def getdocxexcel(self, filename):
        # docx = 'docx'
        #doc转docx格式保存
        # if '.docx' not in filename:
        if filename.endswith('.doc'):
            word = wc.Dispatch('word.application')
            doc = word.Documents.Open(filename)
            docxfilename = filename+'x'
            doc.SaveAs(docxfilename, 12)
            word.Application.Quit()
            d = Document(docxfilename)
            excelfilename = (filename.strip('doc'))+'xls'
        else:
            d = Document(filename)
            excelfilename = (filename.strip('docx'))+'xls'
        #过滤文本
        declude = ['', '长期医嘱：', '临时医嘱：', '出院医嘱：', '长期医嘱:', '临时医嘱:', '出院医嘱:']
        misorder = ['主','要','诊','疗','工','作','重','点','医','嘱']        
        #总行号
        linenumber = 0
        wbk = xlwt.Workbook(encoding='utf-8', style_compression=0)
        #表格建立，首行写入标题
        sheet = wbk.add_sheet('数据', cell_overwrite_ok=True)
        sheet.write(0, 0, '项目名称')
        sheet.write(0, 1, '归属路径')
        sheet.write(0, 2, '单元标记')
        excellinenumber = 1
        #取每一个表格，检查表格是否为7行，表格行数存tablelength中
        tablelength = []
        for t in d.tables:
            tablelength.append(len(t.rows))
        #print('debug:tablelength=', tablelength)
        #标准表格为7行
        for i in tablelength:
            if not i == 7:
                return 1
        #取每一个表格
        for t in d.tables:
            #取每一列
            for columnnumber, columnelement in enumerate(t.columns):
                #第一列不需要
                if columnnumber>0:
                    # 取每一大格
                    linenumber += 1
                    for j,cellelement in enumerate(columnelement.cells):
                        #取第二、三、四行
                        if j > 0 and j < 4:
                            #生成单元标记
                            danyuanbiaoji = str(linenumber) + '-' + str(j)
                            #文本内容分行
                            textindex = cellelement.text.splitlines()
                            for line in textindex:
                                if line in misorder:
                                    return 1                                
                                #去掉文本中的'**医嘱：'以及空行
                                if line not in declude:
                                    #去掉文本中的‘□’
                                    if '□' in line:
                                        newline = line.strip()
                                        newline = newline[1:]
                                        newline = newline.strip()
                                        # print(newline, danyuanbiaoji)
                                        sheet.write(excellinenumber, 0, newline)
                                    else:
                                        line = line.strip()
                                        # print(line, danyuanbiaoji)
                                        sheet.write(excellinenumber, 0, line)
                                    sheet.write(excellinenumber, 2, danyuanbiaoji)
                                    excellinenumber += 1
        wbk.save(excelfilename)
        return 0


test = Solution()
# test.getdirfiles('D:\\WORK-rmyy\\临床路径')
test.getdocxexcel('D:\\WORK-rmyy\\临床路径\\2009年已发布的临床路径(（899-1010）\\自然临产阴道分.doc')
#非标准格式文件：上睑下垂.doc
#test.getdocxexcel('D:\\WORK-rmyy\\临床路径\\2009年已发布的临床路径(（899-1010）\\上睑下垂.doc')
